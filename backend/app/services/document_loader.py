import logging
import mimetypes
from io import BytesIO
from pathlib import Path

import fitz
import pandas as pd
import textract
from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    '.txt',
    '.md',
    '.pdf',
    '.doc',
    '.docx',
    '.csv',
    '.png',
    '.jpg',
    '.jpeg',
    '.webp',
}

IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.webp'}


def _image_suffix_from_content_type(content_type: str | None) -> str:
    guessed = mimetypes.guess_extension(content_type or '')
    if guessed in IMAGE_EXTENSIONS:
        return guessed
    return '.png'


def _should_keep_extracted_image(image_bytes: bytes, *, min_dimension: int, min_bytes: int) -> bool:
    if not image_bytes or len(image_bytes) < min_bytes:
        return False
    try:
        with Image.open(BytesIO(image_bytes)) as image:
            width, height = image.size
            return width >= min_dimension and height >= min_dimension
    except Exception:
        return False

def load_document(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    logger.info('Loading document path=%s ext=%s', file_path, ext)

    if ext not in SUPPORTED_EXTENSIONS:
        logger.warning('Unsupported document extension path=%s ext=%s', file_path, ext)
        raise ValueError(f'Unsupported file extension: {ext}')

    if ext in {'.txt', '.md'}:
        text = file_path.read_text(encoding='utf-8', errors='ignore')
        logger.info('Loaded text/markdown chars=%s', len(text))
        return text

    if ext == '.pdf':
        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or '' for page in reader.pages]
        text = '\n'.join(pages)
        logger.info('Loaded pdf pages=%s chars=%s', len(reader.pages), len(text))
        return text

    if ext == '.docx':
        doc = DocxDocument(str(file_path))
        text = '\n'.join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
        logger.info('Loaded docx paragraphs=%s chars=%s', len(doc.paragraphs), len(text))
        return text

    if ext == '.doc':
        raw = textract.process(str(file_path))
        text = raw.decode('utf-8', errors='ignore')
        logger.info('Loaded doc chars=%s', len(text))
        return text

    if ext == '.csv':
        df = pd.read_csv(file_path)
        text = df.to_csv(index=False)
        logger.info('Loaded csv rows=%s cols=%s chars=%s', len(df.index), len(df.columns), len(text))
        return text

    return ''


def extract_document_images(
    file_path: Path,
    output_dir: Path,
    *,
    min_dimension: int = 32,
    min_bytes: int = 2048,
) -> list[dict]:
    ext = file_path.suffix.lower()
    extracted: list[dict] = []
    skipped_small = 0
    output_dir.mkdir(parents=True, exist_ok=True)

    if ext == '.pdf':
        try:
            pdf = fitz.open(str(file_path))
        except Exception as exc:
            logger.warning('fitz could not open pdf file=%s error=%s', file_path.name, exc)
            return extracted
        for page_no, page in enumerate(pdf, start=1):
            try:
                image_list = page.get_images(full=True)
            except Exception as exc:
                logger.warning('Skipping images on pdf page file=%s page=%s error=%s', file_path.name, page_no, exc)
                continue
            for index, img in enumerate(image_list, start=1):
                try:
                    xref = img[0]
                    base_image = pdf.extract_image(xref)
                    image_bytes = base_image.get('image')
                    if not image_bytes:
                        continue
                    if not _should_keep_extracted_image(
                        image_bytes,
                        min_dimension=min_dimension,
                        min_bytes=min_bytes,
                    ):
                        skipped_small += 1
                        continue
                    suffix = f".{base_image.get('ext', 'png')}"
                    if suffix not in IMAGE_EXTENSIONS:
                        suffix = '.png'
                    image_path = output_dir / f'{file_path.stem}_page_{page_no}_image_{index}{suffix}'
                    image_path.write_bytes(image_bytes)
                    extracted.append({'path': image_path, 'page_no': page_no})
                except Exception as exc:
                    logger.warning(
                        'Skipping problematic pdf image file=%s page=%s image_index=%s error=%s',
                        file_path.name, page_no, index, exc,
                    )
        pdf.close()
        logger.info(
            'Extracted pdf images file=%s images=%s skipped_small=%s',
            file_path.name, len(extracted), skipped_small,
        )
        return extracted

    if ext == '.docx':
        doc = DocxDocument(str(file_path))
        seen_partnames: set[str] = set()
        image_index = 0
        for rel in doc.part.rels.values():
            target_part = getattr(rel, 'target_part', None)
            if target_part is None:
                continue
            content_type = getattr(target_part, 'content_type', '') or ''
            if not content_type.startswith('image/'):
                continue
            partname = str(getattr(target_part, 'partname', ''))
            if partname in seen_partnames:
                continue
            seen_partnames.add(partname)
            if not _should_keep_extracted_image(
                target_part.blob,
                min_dimension=min_dimension,
                min_bytes=min_bytes,
            ):
                skipped_small += 1
                continue
            image_index += 1
            suffix = _image_suffix_from_content_type(content_type)
            image_path = output_dir / f'{file_path.stem}_image_{image_index}{suffix}'
            image_path.write_bytes(target_part.blob)
            extracted.append({'path': image_path, 'page_no': None})
        logger.info(
            'Extracted docx images file=%s images=%s skipped_small=%s',
            file_path.name, len(extracted), skipped_small,
        )
        return extracted

    return extracted
